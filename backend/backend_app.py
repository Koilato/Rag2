from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import pandas as pd
import io
from docx import Document
from io import BytesIO

# 导入 TEST.py 中使用的模块
import CsvDataOp
import MySqlSource
import VectorDatabase
import utils
import ALiYunConnection
import json

app = FastAPI()

# 配置CORS，允许Vue前端访问
origins = [
    "http://localhost",
    "http://localhost:8080",  # 您的Vue前端可能运行的地址
    "http://127.0.0.1:8080",
    # 如果您的前端部署在其他域名，请在此处添加
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 全局数据库连接和ChromaDB客户端/集合
# 在实际生产环境中，这些连接应该通过依赖注入或更健壮的方式管理
mysql_connect = None
chroma_client = None


@app.on_event("startup")
async def startup_event():
    """
    应用启动时连接数据库和ChromaDB
    """
    global mysql_connect, chroma_client
    try:
        mysql_connect = MySqlSource.connect_to_mysql()
        chroma_client = VectorDatabase.connect_to_chromadb()
        print("数据库和ChromaDB客户端初始化成功！")
    except Exception as e:
        print(f"数据库或ChromaDB连接失败: {e}")
        # 在生产环境中，这里可能需要更优雅的错误处理，例如退出应用或记录日志

@app.on_event("shutdown")
async def shutdown_event():
    """
    应用关闭时关闭数据库连接
    """
    global mysql_connect
    if mysql_connect:
        mysql_connect.close()
        print("MySQL连接已关闭。")

# 定义请求体模型
class LoginRequest(BaseModel):
    username: str
    password: str

class ProcessDataRequest(BaseModel):
    db_name: str  # 数据库名称，设为必需
    collection_name: str  # 集合名称，设为必需
    cve: Optional[List[str]] = None
    uuid: Optional[List[str]] = None
    host: Optional[List[str]] = None
    plugin_id: Optional[List[str]] = None
    risk: Optional[List[str]] = None
    protocol: Optional[List[str]] = None
    start_date: Optional[str] = None  # YYYY-MM-DD
    end_date: Optional[str] = None  # YYYY-MM-DD
    auto_export_report: Optional[bool] = False
    web_search: Optional[bool] = False
    chat_message: Optional[str] = None
    n_results: Optional[int] = 30  # 默认值

@app.post("/api/export_report")
async def export_report(request: ProcessDataRequest):
    """
    根据前端提供的参数生成并导出.doc格式的报告。
    """
    try:
        document = Document()
        document.add_heading('报告', level=1)

        document.add_paragraph(f"生成日期: {utils.get_today_date_formatted()}")
        document.add_paragraph("\n--- 参数配置 ---")
        if request.cve:
            document.add_paragraph(f"CVE: {', '.join(request.cve)}")
        if request.uuid:
            document.add_paragraph(f"UUID: {', '.join(request.uuid)}")
        if request.host:
            document.add_paragraph(f"HOST: {', '.join(request.host)}")
        if request.plugin_id:
            document.add_paragraph(f"Plugin ID: {', '.join(request.plugin_id)}")
        if request.risk:
            document.add_paragraph(f"Risk: {', '.join(request.risk)}")
        if request.protocol:
            document.add_paragraph(f"Protocol: {', '.join(request.protocol)}")
        if request.start_date and request.end_date:
            document.add_paragraph(f"时间范围: {request.start_date} 至 {request.end_date}")

        document.add_paragraph("\n--- 工具设置 ---")
        document.add_paragraph(f"自动导出报告: {'启用' if request.auto_export_report else '禁用'}")
        document.add_paragraph(f"联网搜索: {'启用' if request.web_search else '禁用'}")

        # 如果有聊天消息，也可以包含在报告中
        if request.chat_message:
            document.add_paragraph("\n--- 会话内容 ---")
            document.add_paragraph(f"用户消息: {request.chat_message}")

        # 这里可以根据需要添加更多从LLM获取的数据或数据库查询结果
        # 例如，可以调用 process_data 接口的逻辑来获取最终答案，并将其添加到报告中
        # 为了简化，这里只包含参数信息

        # 将文档保存到内存中
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0) # 将文件指针移到开头

        # 返回文件作为响应
        filename = f"报告_{utils.get_today_date_formatted()}.docx"
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成报告失败: {e}")

@app.post("/api/login")
async def login(request: LoginRequest):
    try:
        with open("backend/users.json", "r") as f:
            users = json.load(f)

        if request.username in users and users[request.username] == request.password:
            # 验证成功，创建数据库和集合
            db_name = request.username
            collection_name = request.username

            # 创建 MySQL 数据库 (如果不存在)
            MySqlSource.create_database(mysql_connect, db_name)

            # 创建 ChromaDB 集合 (如果不存在)
            chroma_client.get_or_create_collection(name=collection_name)
            
            print(f"为用户 '{db_name}' 确保数据库和集合已就绪。")

            return JSONResponse(content={"success": True, "message": "Login successful"})
        else:
            return JSONResponse(status_code=401, content={"success": False, "message": "Invalid username or password"})
    except FileNotFoundError:
        raise HTTPException(status_code=500, detail="Users file not found.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {e}")

@app.post("/api/upload_csv")
async def upload_csv(
        db_name: str = Form(...),
        table_name: str = Form(...),
        collection_name: str = Form(...),
        file: UploadFile = File(...)
):
    """
    Receives an uploaded CSV file, processes it, and stores it in the specified
    MySQL database/table and ChromaDB collection.
    """
    if not file.filename.endswith('.csv'):
        raise HTTPException(status_code=400, detail="Only CSV files are accepted.")

    try:
        contents = await file.read()
        sio = io.StringIO(contents.decode('utf-8'))
        pdfile = CsvDataOp.transform_nessus_data(sio)

        if pdfile.empty:
            raise HTTPException(status_code=400, detail="Processed CSV file is empty.")

        # Create database and table
        MySqlSource.create_database(mysql_connect, db_name)
        MySqlSource.create_pd_table(mysql_connect, db_name, table_name, use_vulnerability_template=True)

        # Add IDs and UUIDs
        pdfile = utils.add_ids_and_uuid(mysql_connect, db_name, pdfile, table_name)

        # Insert data into MySQL
        MySqlSource.insert_vulnerability_data(mysql_connect, db_name, table_name, pdfile)

        # Sync data from MySQL to ChromaDB
        MySqlSource.sync_mysql_to_chromadb(
            connection=mysql_connect,
            chroma_client=chroma_client,
            db_name=db_name,
            table_name=table_name,
            collection_name=collection_name,
            batch_size=1000
        )

        return JSONResponse(
            content={
                "message": f"File {file.filename} processed and stored successfully!",
                "database": db_name,
                "table": table_name,
                "collection": collection_name
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File processing failed: {e}")

@app.post("/api/process_data")
async def process_data(request: ProcessDataRequest):
    """
    接收前端输入的文本参数和会话消息，执行LLM查询并返回结果。
    """
    if not mysql_connect or not chroma_client:
        raise HTTPException(status_code=503, detail="后端服务未完全启动或数据库连接失败。")
    # 获取指定的ChromaDB集合
    try:
        chroma_collection = chroma_client.get_or_create_collection(name=request.collection_name)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"无法获取ChromaDB集合'{request.collection_name}': {e}")
    question_parts = []
    if request.cve:
        question_parts.append(f"CVE: {', '.join(request.cve)}")
    if request.uuid:
        question_parts.append(f"UUID: {', '.join(request.uuid)}")
    if request.host:
        question_parts.append(f"HOST: {', '.join(request.host)}")
    if request.plugin_id:
        question_parts.append(f"Plugin ID: {', '.join(request.plugin_id)}")
    if request.risk:
        question_parts.append(f"Risk: {', '.join(request.risk)}")
    if request.protocol:
        question_parts.append(f"Protocol: {', '.join(request.protocol)}")
    if request.start_date and request.end_date:
        question_parts.append(f"时间范围从 {request.start_date} 到 {request.end_date}")
    if request.chat_message:
        question_parts.append(request.chat_message)

    if not question_parts:
        raise HTTPException(status_code=400, detail="请提供至少一个查询参数或会话消息。")

    question = " ".join(question_parts)
    if request.auto_export_report:
        question += " 自动导出报告。"
    if request.web_search:
        question += " 进行联网搜索。"

    try:
        # LLM查询流程 (从 TEST.py 提取并适应)
        file_path_json_prompt = "prompt_for_json.txt"
        with open(file_path_json_prompt, 'r', encoding='utf-8') as file:
            prompt1 = file.read()

        message1 = [
            {"role": "system", "content": f"{prompt1}"},
            {"role": "user", "content": f"{question}"}
        ]

        response1 = ALiYunConnection.qwen_query(message1, stream=False, enable_thinking=False)
        full_text1 = response1.choices[0].message.content
        parsed_json_output = utils.parse_llm_json_output(full_text1)

        where_clause = parsed_json_output.get("chroma_where_filter", {})
        query_text = parsed_json_output.get("query_text", question) # 如果LLM没有提供query_text，则使用原始问题

        uuids = VectorDatabase.query_vulnerabilities_for_uuids(query_text, request.n_results, where_clause, chroma_collection)

        docs = []
        if uuids:
            # 使用请求中提供的db_name
            docs = VectorDatabase.get_full_documents_from_mysql(uuids, db_name=request.db_name)
            docs = utils.format_mysql_dataframe_for_llm(docs)
        else:
            docs = "未找到相关文档。"

        today_time = utils.get_today_date_formatted()
        file_path_answer_prompt = "prompt_for_answer.txt"
        with open(file_path_answer_prompt, 'r', encoding='utf-8') as file:
            prompt2 = file.read()

        new_message = [
            {"role": "system", "content": f"{prompt2}"},
            {"role": "user", "content": f"{question}：今天的日期是：{today_time}以下是相关文本：\n\n{docs},要输出成markdown表格的形式，同时每一回答来源都要给出原始的uuid" }
        ]

        response2 = ALiYunConnection.qwen_query(new_message, stream=False, enable_thinking=False)
        final_answer = response2.choices[0].message.content

        return JSONResponse(content={"answer": final_answer})

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"处理请求失败: {e}")

# 运行FastAPI应用 (在开发环境中，通常通过命令行 `uvicorn backend_app:app --reload` 运行)
# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)
